# -*- coding: utf-8 -*-
"""
简历解析模块 v2 - AI增强版
================================
升级内容（v2）：
  1. 新增 OCR 支持：图片简历（jpg/png）→ 文本
  2. 新增 AI 深度解析：LLM 提取完整结构化 JSON
  3. 新增全流程管线：文件字节 → 文本 → AI解析 → resume_parsed JSON
  4. 向后兼容：原有 parse_resume() 和提取函数全部保留

依赖：
  - 基础（已有）: pdfplumber, python-docx
  - OCR（按需）: paddleocr（可选，未安装时图片简历会报错）
  - AI（按需）: openai 库（调用 LLM API）

环境变量：
  OPENAI_API_KEY  - LLM API Key
  OPENAI_API_BASE - LLM API Base URL
  AI_PARSE_MODEL  - 模型名（默认 gpt-4o-mini）
"""

import re
import os
import io
import json
import base64
import traceback
from typing import Optional, Tuple


# ============================================================
# 配置
# ============================================================
AI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
AI_API_BASE = os.environ.get("OPENAI_API_BASE", "")
AI_MODEL = os.environ.get("AI_PARSE_MODEL", "deepseek-chat")


# ============================================================
# 第1层：文本提取（PDF / DOCX / TXT / 图片）
# ============================================================

def extract_text(file_bytes: bytes, filename: str) -> str:
    """从上传的简历文件中提取纯文本。支持 PDF、DOCX、TXT、JPG/PNG。"""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"):
        return _extract_image(file_bytes, ext)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，请上传 PDF、DOCX、TXT 或图片文件")


def _extract_pdf(file_bytes: bytes) -> str:
    """从 PDF 提取文本"""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    """从 DOCX 提取文本（含表格内容）"""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # 提取表格内容（常见于格式化的简历）
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def _extract_image(file_bytes: bytes, ext: str) -> str:
    """从图片简历中提取文本（OCR）。

    优先使用 PaddleOCR（中文识别好），fallback 到 Tesseract。
    如果两者都不可用，抛出友好错误。
    """
    # 尝试 PaddleOCR
    try:
        from paddleocr import PaddleOCR
        import numpy as np
        from PIL import Image

        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        image = Image.open(io.BytesIO(file_bytes))

        # 转 RGB（处理 RGBA/灰度）
        if image.mode != "RGB":
            image = image.convert("RGB")
        img_array = np.array(image)

        result = ocr.ocr(img_array, cls=True)
        if not result or not result[0]:
            raise Exception("PaddleOCR returned empty result")

        lines = []
        for line_info in result[0]:
            text = line_info[1][0]
            lines.append(text)
        return "\n".join(lines)

    except ImportError:
        pass  # PaddleOCR not installed, try tesseract

    # 尝试 Tesseract
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        if text.strip():
            return text.strip()
    except ImportError:
        pass

    raise ImportError(
        "图片简历需要 OCR 引擎。请安装其中之一：\n"
        "  PaddleOCR: pip install paddlepaddle paddleocr\n"
        "  Tesseract: pip install pytesseract (需额外安装 tesseract-OCR)"
    )


# ============================================================
# 第2层：正则解析（原有，快速本地回退）
# ============================================================

def parse_resume(text: str) -> dict:
    """
    从简历文本中解析结构化信息（正则版）。
    返回字段: name, phone, email, skills, work_years, recent_position,
              recent_company, resume_text, education
    """
    if not text:
        return {}

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    full_text = "\n".join(lines)

    result = {
        "name": _extract_name(lines, full_text),
        "phone": _extract_phone(full_text),
        "email": _extract_email(full_text),
        "skills": _extract_skills(full_text),
        "work_years": _extract_work_years(full_text),
        "recent_position": _extract_recent_position(lines),
        "recent_company": _extract_recent_company(lines, full_text),
        "education": _extract_education(full_text),
        "resume_text": full_text,
    }

    return result


# ===== 原有正则提取函数（保持不变）=====

def _extract_name(lines: list, text: str) -> str:
    """提取姓名 - 取第一行（通常是姓名）"""
    if not lines:
        return ""
    first_line = lines[0].strip()
    for prefix in ["个人简历", "简历", "RESUME", "CURRICULUM VITAE", "CV"]:
        if first_line.upper().startswith(prefix):
            if len(lines) > 1:
                return lines[1].strip()[:20]
            return ""
    if len(first_line) <= 20:
        return first_line
    m = re.search(r"姓\s*名[：:]\s*(.{2,10})", text)
    if m:
        return m.group(1).strip()
    return first_line[:20]


def _extract_phone(text: str) -> str:
    """提取手机号"""
    m = re.search(r"1[3-9]\d{9}", text)
    if m:
        return m.group(0)
    m = re.search(r"(\d{3,4})[-\s]?(\d{7,8})", text)
    if m:
        return f"{m.group(1)}-{m.group(2)}"
    return ""


def _extract_email(text: str) -> str:
    """提取邮箱"""
    m = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else ""


def _extract_skills(text: str) -> str:
    """提取技能"""
    patterns = [
        r"(?:专业技能|技术技能|技能|特长|技能特长|掌握|熟悉|精通)[：:\s]*([^\n]{10,200}?)(?:\n|$)",
        r"(?:熟练|掌握|精通|了解)\s*([A-Za-z\s,，、/]+(?:框架|语言|工具|技术|库|系统)?)",
    ]
    for pattern in patterns:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            skill_text = m.group(1).strip()
            skill_text = re.sub(r"[。.!！？?]", "", skill_text)
            if 2 < len(skill_text) < 200:
                return skill_text[:200]

    tech_keywords = re.findall(
        r"\b(Python|Java|JavaScript|TypeScript|React|Vue|Angular|Node\.?js|Go|Rust|"
        r"C\+\+|C#|Swift|Kotlin|PHP|Ruby|SQL|NoSQL|MySQL|PostgreSQL|MongoDB|Redis|"
        r"Docker|Kubernetes|AWS|Azure|Linux|Git|TensorFlow|PyTorch|"
        r"HTML|CSS|Sass|Webpack|Nginx|Spring|Django|Flask|FastAPI|"
        r"Spark|Hadoop|Flink|Kafka|RabbitMQ|Elasticsearch|"
        r"Photoshop|Figma|Sketch|Illustrator|Axure|"
        r"Excel|PPT|Word|Tableau|Power\s*BI)\b",
        text,
        re.IGNORECASE,
    )
    if tech_keywords:
        seen = set()
        unique = []
        for kw in tech_keywords:
            if kw.lower() not in seen:
                seen.add(kw.lower())
                unique.append(kw)
        return ", ".join(unique)
    return ""


def _extract_work_years(text: str) -> int:
    """提取工作年限"""
    patterns = [
        r"(\d+)\s*年.*?经验",
        r"(\d+)\s*年(?:工作)?经验",
        r"工作(?:年限|经历)[：:]\s*(\d+)\s*年",
        r"从业\s*(\d+)\s*年",
        r"(?:毕业|离校)\s*(\d+)\s*年",
    ]
    for pattern in patterns:
        m = re.search(pattern, text)
        if m:
            try:
                years = int(m.group(1))
                if 0 < years <= 50:
                    return years
            except ValueError:
                pass
    return 0


def _extract_recent_position(lines: list) -> str:
    """提取最近职位"""
    work_section = False
    for i, line in enumerate(lines):
        if re.search(r"(?:工作经历|工作(经验|经历)|职业经历|工作背景|实习经历)", line):
            work_section = True
            continue
        if work_section:
            if not line or re.match(r"[-=_*#]{3,}", line):
                continue
            if re.search(r"(?:公司|科技|集团|有限|股份|Co\.|Inc\.|Ltd\.)", line):
                continue
            if re.match(r"\d{4}[\.\-/年]\s*\d{0,2}", line):
                continue
            if 2 < len(line) <= 30:
                return line.strip()
    return ""


def _extract_recent_company(lines: list, text: str) -> str:
    """提取最近公司"""
    for line in lines:
        m = re.search(r"(?:公司|科技|集团|有限|股份)\s*[）)]?\s*$", line)
        if m and 4 < len(line) <= 50:
            company = line.strip()
            company = re.sub(
                r"^\d{4}[\.\-/年]\s*\d{0,2}[\.\-/月]?\s*[-–~—至到]\s*\d{4}[\.\-/年]?\s*\d{0,2}?\s*",
                "", company
            )
            if company.strip():
                return company.strip()[:50]
    return ""


def _extract_education(text: str) -> str:
    """提取学历信息"""
    edu_keywords = {
        "博士": "博士", "硕士": "硕士", "研究生": "硕士",
        "本科": "本科", "学士": "本科", "大专": "大专",
        "专科": "大专", "高中": "高中", "专升本": "本科",
        "MBA": "MBA", "EMBA": "EMBA",
    }
    for keyword, level in edu_keywords.items():
        if keyword in text:
            return level
    m = re.search(r"([\u4e00-\u9fa5]{2,15}(?:大学|学院|学校|研究院|研究所))", text)
    if m:
        return m.group(1)
    return ""


# ============================================================
# 第3层：AI 深度解析（新增 v2）
# ============================================================

# ─── AI Prompt 模板 ───

RESUME_PARSE_PROMPT = """你是一位资深招聘专家。请仔细阅读以下简历内容，提取所有关键信息，
并严格按 JSON 格式返回。不要遗漏任何信息。

## 输出格式要求（严格遵守）

{
  "name": "姓名（字符串）",
  "phone": "手机号（字符串，没有则为空）",
  "email": "邮箱（字符串，没有则为空）",
  "gender": "性别（男/女/未知）",
  "age": 年龄数字或null,
  "current_location": "现居城市（如'北京'，没有则为空）",
  "education": [
    {
      "school": "学校全称",
      "degree": "学历（博士/硕士/本科/大专/高中等）",
      "major": "专业名称",
      "start": "起始年份（如'2016'）",
      "end": "结束年份（如'2020'，在读则填'至今'）"
    }
  ],
  "work_experience": [
    {
      "company": "公司全称",
      "position": "职位名称",
      "start": "起始时间（如'2019-07'）",
      "end": "结束时间（如'2023-03'，在职则填'至今'）",
      "description": "工作描述摘要（50-200字）",
      "industry": "行业（如'互联网'、'金融'）"
    }
  ],
  "project_experience": [
    {
      "name": "项目名称",
      "role": "担任角色",
      "description": "项目描述（50-150字）",
      "tech_stack": ["使用的技术"]
    }
  ],
  "skills": ["技能1", "技能2"],
  "certifications": ["证书1", "证书2"],
  "languages": ["语言1（如'英语:CET-6'）"],
  "current_salary": "当前薪资（如'30K-40K'或'月薪3万'，没有则为空）",
  "expected_salary": "期望薪资（没有则为空）",
  "work_years": 工作年限数字或null,
  "job_hopping_frequency": "跳槽频率评估（低/中/高，'低'表示稳定）",
  "self_assessment": "自我评价（保留原文，最多300字）",
  "raw_summary": "简历整体摘要（2-3句话概括候选人画像）"
}

## 提取规则

1. **教育经历**：按时间倒序排列。如果简历中只有学校名没有专业，专业留空。
2. **工作经历**：按时间倒序排列。每段经历提取公司、职位、时间、职责关键词。
3. **技能**：列出所有提到的硬技能和软技能，不限于技术（含管理、沟通等）。
4. **证书**：列出所有证书、资格证、获奖。
5. **薪资**：区分当前薪资和期望薪资，注意"年薪"vs"月薪"，统一标注。
6. **工作年限**：优先取简历中明确标注的年限；否则计算最早工作经历至今。
7. **跳槽频率**：每段工作<1年且>=3段 → 高；1-2年 → 中；>2年 → 低。
8. **缺失字段填空字符串或null（数字字段）**，不要编造。

## 简历内容

{resume_text}

## 重要要求

1. **不要截断任何内容**，所有字段必须完整提取
2. **工作经历和项目经历必须全部列出**，不要省略任何一段
3. **如果简历内容很多，确保输出完整**，不要因为长度限制而省略字段
4. **只输出 JSON**，不要带任何 markdown 标记或额外说明"""


def _call_llm(prompt: str, api_key: str = None, api_base: str = None,
              model: str = None) -> str:
    """调用 LLM API"""
    from openai import OpenAI

    client = OpenAI(
        api_key=api_key or AI_API_KEY,
        base_url=api_base or AI_API_BASE,
    )

    response = client.chat.completions.create(
        model=model or AI_MODEL,
        messages=[
            {"role": "system", "content": "你是一位资深的招聘专家，擅长从简历中提取结构化信息。只输出JSON，不输出其他内容。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,       # 低温度保证输出稳定
        max_tokens=16384,
    )

    return response.choices[0].message.content


def _clean_json_response(text: str) -> str:
    """清理 LLM 返回的 JSON，去掉 markdown 标记等"""
    text = text.strip()
    # 去掉 ```json ... ``` 包裹
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*\n?", "", text)
        text = re.sub(r"\n?```\s*$", "", text)
    # 处理可能的 HTML 实体
    text = text.replace("&quot;", '"').replace("&#39;", "'").replace("&amp;", "&")
    return text.strip()


def parse_resume_with_ai(text: str, api_key: str = None,
                         api_base: str = None, model: str = None,
                         max_chars: int = 24000) -> dict:
    """
    使用 LLM 深度解析简历文本，返回完整结构化 JSON。

    参数：
        text: 简历纯文本
        api_key: LLM API Key（可选，默认取环境变量）
        api_base: LLM API Base URL（可选）
        model: LLM 模型名（可选）
        max_chars: 最大输入字符数（超长简历会截断，避免 token 超限）

    返回：
        dict: 结构化简历数据，格式见 RESUME_PARSE_PROMPT 模板
    """
    if not text or not text.strip():
        return {}

    # 检查 API Key
    key = api_key or AI_API_KEY
    if not key:
        raise ValueError(
            "AI 解析需要 LLM API Key。请设置环境变量 OPENAI_API_KEY，"
            "或传入 api_key 参数。也可以改用 parse_resume() 进行正则解析。"
        )

    # 截断过长文本（保留开头 + 结尾，中间省略）
    if len(text) > max_chars:
        half = max_chars // 2
        text = text[:half] + "\n\n...（中间部分省略）...\n\n" + text[-half:]

    # 用 replace 而非 format，避免模板中 JSON 示例的 {} 被 Python 误解析为占位符
    prompt = RESUME_PARSE_PROMPT.replace("{resume_text}", text)

    try:
        raw_response = _call_llm(prompt, key, api_base, model)
        cleaned = _clean_json_response(raw_response)
        parsed = json.loads(cleaned)

        # 基本校验
        if not isinstance(parsed, dict):
            raise ValueError(f"AI 返回的不是 JSON 对象: {type(parsed)}")

        return parsed
    except json.JSONDecodeError as e:
        raise ValueError(
            f"AI 返回的内容无法解析为 JSON:\n{str(e)[:200]}\n\n"
            f"原始返回内容（前500字符）:\n{raw_response[:500]}"
        )
    except Exception as e:
        raise RuntimeError(f"AI 解析失败: {e}")


# ============================================================
# 第4层：全流程管线（新增 v2）
# ============================================================

def parse_resume_file(file_bytes: bytes, filename: str,
                      use_ai: bool = True,
                      api_key: str = None,
                      api_base: str = None,
                      model: str = None) -> dict:
    """
    简历文件 → 结构化 JSON 的完整管线。

    流程：
      1. 提取文本（PDF/DOCX/TXT/图片）
      2. 正则快速提取基本信息
      3. [可选] AI 深度解析生成完整 JSON
      4. 合并结果，优先 AI 结果，缺失字段用正则补充

    参数：
        file_bytes: 文件字节内容
        filename: 文件名（用于判断格式）
        use_ai: 是否使用 AI 深度解析（默认 True）
        api_key, api_base, model: 传给 parse_resume_with_ai

    返回：
        dict: {
            "text": "简历纯文本",
            "basic": {...},          # 正则解析结果
            "ai_parsed": {...},      # AI 解析结果（仅 use_ai=True 时）
            "merged": {...},         # 合并后的最终结果
            "parser_method": "ai" | "regex",
            "parser_version": "v2"
        }
    """
    result = {
        "text": "",
        "basic": {},
        "ai_parsed": None,
        "merged": {},
        "parser_method": "regex",
        "parser_version": "v2",
        "error": None,
    }

    try:
        # Step 1: 提取文本
        text = extract_text(file_bytes, filename)
        if not text.strip():
            raise ValueError("未能从文件中提取到任何文本内容")
        result["text"] = text

        # Step 2: 正则快速解析
        basic = parse_resume(text)
        result["basic"] = basic

        # Step 3: AI 深度解析
        if use_ai:
            try:
                ai_parsed = parse_resume_with_ai(
                    text, api_key=api_key, api_base=api_base, model=model
                )
                result["ai_parsed"] = ai_parsed
                result["parser_method"] = "ai"
            except (ValueError, RuntimeError) as e:
                # AI 失败时回退到正则
                result["error"] = str(e)

        # Step 4: 合并结果
        result["merged"] = _merge_results(basic, result.get("ai_parsed"))

    except Exception as e:
        result["error"] = str(e)
        traceback.print_exc()

    return result


def _merge_results(basic: dict, ai_parsed: Optional[dict]) -> dict:
    """
    合并正则解析和 AI 解析结果。
    优先取 AI 结果，缺失字段用正则补充。
    """
    if ai_parsed is None:
        # 纯正则模式：转换格式为结构化标准格式
        return _basic_to_structured(basic)

    merged = dict(ai_parsed)  # 从 AI 结果开始

    # 用正则结果补充缺失字段
    if not merged.get("name") and basic.get("name"):
        merged["name"] = basic["name"]
    if not merged.get("phone") and basic.get("phone"):
        merged["phone"] = basic["phone"]
    if not merged.get("email") and basic.get("email"):
        merged["email"] = basic["email"]
    if not merged.get("work_years") and basic.get("work_years"):
        merged["work_years"] = basic["work_years"]
    if not merged.get("skills") and basic.get("skills"):
        merged["skills"] = [basic["skills"]]
    if not merged.get("education") and basic.get("education"):
        merged["education"] = [{"school": basic["education"], "degree": "", "major": "", "start": "", "end": ""}]

    # 补充来源标记
    merged["_parser_method"] = "ai" if ai_parsed else "regex"
    merged["_parser_version"] = "v2"

    return merged


def _basic_to_structured(basic: dict) -> dict:
    """将正则解析结果转换为结构化格式（与 AI 输出统一）"""
    result = {
        "name": basic.get("name", ""),
        "phone": basic.get("phone", ""),
        "email": basic.get("email", ""),
        "work_years": basic.get("work_years", 0),
        "skills": [basic["skills"]] if basic.get("skills") else [],
        "education": [],
        "work_experience": [],
        "self_assessment": "",
        "_parser_method": "regex",
        "parser_method": "regex",
        "_parser_version": "v2",
    }

    if basic.get("education"):
        result["education"] = [{
            "school": basic["education"],
            "degree": "",
            "major": "",
            "start": "",
            "end": "",
        }]

    if basic.get("recent_position") or basic.get("recent_company"):
        result["work_experience"] = [{
            "company": basic.get("recent_company", ""),
            "position": basic.get("recent_position", ""),
            "start": "", "end": "",
            "description": "",
            "industry": "",
        }]

    return result


# ============================================================
# 输出辅助
# ============================================================

def summarize_for_display(parsed: dict or object) -> str:
    """
    生成解析结果摘要，用于界面展示。
    兼容旧版 dict 和新版 parse_resume_file() 返回的结构。
    """
    # 检测是新版还是旧版格式
    if isinstance(parsed, dict) and "merged" in parsed:
        # 新版管线输出，使用 merged 字段
        data = parsed["merged"]
        method = parsed.get("parser_method", "regex")
    else:
        data = parsed
        method = "regex"

    parts = []

    # 解析方式
    tag = "[AI]" if method == "ai" else "[Regex]"
    parts.append(f"解析方式: {tag}")

    if data.get("name"):
        parts.append(f"姓名: {data['name']}")
    if data.get("phone"):
        parts.append(f"电话: {data['phone']}")
    if data.get("email"):
        parts.append(f"邮箱: {data['email']}")
    if data.get("gender"):
        parts.append(f"性别: {data['gender']}")
    if data.get("current_location"):
        parts.append(f"所在地: {data['current_location']}")
    if data.get("work_years"):
        parts.append(f"工作年限: {data['work_years']}年")

    # 学历（新版格式）
    education = data.get("education", [])
    if isinstance(education, list) and education:
        edu_strs = []
        for edu in education[:2]:  # 最多显示两个
            if isinstance(edu, dict):
                parts_list = []
                if edu.get("school"): parts_list.append(edu["school"])
                if edu.get("degree"): parts_list.append(edu["degree"])
                if edu.get("major"): parts_list.append(edu["major"])
                if parts_list:
                    edu_strs.append(" ".join(parts_list))
        if edu_strs:
            parts.append(f"学历: {'; '.join(edu_strs)}")
    elif isinstance(education, str) and education:
        parts.append(f"学历: {education}")

    # 工作经历（新版格式）
    work_exp = data.get("work_experience", [])
    if isinstance(work_exp, list) and work_exp:
        latest = work_exp[0]
        if isinstance(latest, dict):
            pos = latest.get("position", "")
            comp = latest.get("company", "")
            if pos or comp:
                parts.append(f"最近工作: {pos} @ {comp}")
    elif isinstance(data.get("recent_position"), str) and data["recent_position"]:
        parts.append(f"最近职位: {data['recent_position']}")
    elif isinstance(data.get("recent_company"), str) and data["recent_company"]:
        parts.append(f"最近公司: {data['recent_company']}")

    # 技能
    skills = data.get("skills", [])
    if isinstance(skills, list) and skills:
        parts.append(f"技能: {', '.join(skills[:8])}")
    elif isinstance(skills, str) and skills:
        parts.append(f"技能: {skills[:100]}")

    # 错误信息
    if isinstance(parsed, dict) and parsed.get("error"):
        parts.insert(0, f"[警告] AI解析失败，使用正则回退: {parsed['error'][:100]}")

    # 新版专用：raw_summary
    if method == "ai" and data.get("raw_summary"):
        parts.append(f"\n候选人画像:\n  {data['raw_summary']}")

    return "\n".join(parts) if parts else "未能提取到有效信息"


def parsed_to_db_fields(parsed: dict) -> dict:
    """
    将 parse_resume_file() 的结果映射为 candidates 表的字段。

    返回 dict 可直接传给 db.add_candidate() 或 db.update_candidate()：
      - name, phone, email, work_years, education（覆盖原有逻辑）
      - skills（逗号分隔）
      - recent_position, recent_company（从 work_experience[0] 提取）
      - resume_parsed（完整 JSON 字符串，存入 resume_parsed 列）
    """
    merged = parsed.get("merged", parsed)

    # 提取最近职位和公司（新版格式）
    work_exp = merged.get("work_experience", [])
    recent_position = ""
    recent_company = ""
    if isinstance(work_exp, list) and work_exp:
        first = work_exp[0]
        if isinstance(first, dict):
            recent_position = first.get("position", "")
            recent_company = first.get("company", "")

    # 技能处理：list → 逗号分隔字符串
    skills = merged.get("skills", [])
    if isinstance(skills, list):
        skills_str = ", ".join(skills)
    else:
        skills_str = str(skills) if skills else ""

    # 学历处理
    education = merged.get("education", [])
    if isinstance(education, list) and education:
        edu = education[0]
        if isinstance(edu, dict):
            edu_str = f"{edu.get('degree', '')} - {edu.get('school', '')} - {edu.get('major', '')}".strip(" -")
        else:
            edu_str = str(edu)
    elif isinstance(education, str):
        edu_str = education
    else:
        edu_str = ""

    return {
        "name": merged.get("name", ""),
        "phone": merged.get("phone", ""),
        "email": merged.get("email", ""),
        "work_years": merged.get("work_years", 0),
        "skills": skills_str,
        "education": edu_str,
        "recent_position": recent_position,
        "recent_company": recent_company,
        "resume_parsed": json.dumps(merged, ensure_ascii=False),
        "parser_method": "ai" if parsed.get("ai_parsed") else "regex",
    }